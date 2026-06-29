import os
import re
from pathlib import Path
from dataclasses import dataclass, field
from typing import Optional
import fitz  # PyMuPDF
import docx  # python-docx


@dataclass
class ParsedDocument:
    """
    The result of parsing any document.
    This is the clean, structured object that the rest of the system works with.
    """
    file_name: str
    file_type: str
    raw_text: str
    pages: list[str]
    page_count: int
    detected_title: Optional[str]
    detected_date: Optional[str]
    document_category: str
    char_count: int
    word_count: int


class DocumentParser:
    """
    Parses PDF, DOCX, and plain text files into clean ParsedDocument objects.
    """

    SUPPORTED_TYPES = {".pdf", ".docx", ".txt", ".md"}

    def parse(self, file_path: str) -> ParsedDocument:
        """
        Main entry point. Detect file type and route to the right parser.
        """
        path = Path(file_path)

        if not path.exists():
            raise FileNotFoundError(f"File not found: {file_path}")

        if path.suffix.lower() not in self.SUPPORTED_TYPES:
            raise ValueError(
                f"Unsupported file type: {path.suffix}. "
                f"Supported: {self.SUPPORTED_TYPES}"
            )

        if path.suffix.lower() == ".pdf":
            pages, raw_text = self._parse_pdf(path)
        elif path.suffix.lower() == ".docx":
            pages, raw_text = self._parse_docx(path)
        else:
            pages, raw_text = self._parse_text(path)

        cleaned_text = self._clean_text(raw_text)

        return ParsedDocument(
            file_name=path.name,
            file_type=path.suffix.lower(),
            raw_text=cleaned_text,
            pages=pages,
            page_count=len(pages),
            detected_title=self._detect_title(cleaned_text, path.name),
            detected_date=self._detect_date(cleaned_text),
            document_category=self._detect_category(cleaned_text, path.name),
            char_count=len(cleaned_text),
            word_count=len(cleaned_text.split()),
        )

    def _parse_pdf(self, path: Path) -> tuple[list[str], str]:
        """
        Extract text from a PDF file, page by page.
        Uses PyMuPDF (fitz) for high-accuracy extraction.
        """
        pages = []
        doc = fitz.open(str(path))

        for page_num in range(len(doc)):
            page = doc[page_num]
            page_text = page.get_text("text")
            pages.append(page_text)

        doc.close()
        raw_text = "\n\n".join(pages)
        return pages, raw_text

    def _parse_docx(self, path: Path) -> tuple[list[str], str]:
        """
        Extract text from a DOCX file.
        DOCX has no concept of pages, so we treat each paragraph as a unit.
        Groups every 20 paragraphs into a synthetic 'page' for consistency.
        """
        document = docx.Document(str(path))
        paragraphs = []

        for paragraph in document.paragraphs:
            text = paragraph.text.strip()
            if text:
                paragraphs.append(text)

        chunk_size = 20
        pages = []
        for i in range(0, len(paragraphs), chunk_size):
            page_paragraphs = paragraphs[i: i + chunk_size]
            pages.append("\n\n".join(page_paragraphs))

        raw_text = "\n\n".join(paragraphs)
        return pages, raw_text

    def _parse_text(self, path: Path) -> tuple[list[str], str]:
        """
        Read plain text or markdown files.
        Splits on double newlines to create synthetic page breaks.
        """
        with open(path, "r", encoding="utf-8", errors="replace") as f:
            raw_text = f.read()

        sections = [s.strip() for s in raw_text.split("\n\n") if s.strip()]
        chunk_size = 10
        pages = []
        for i in range(0, len(sections), chunk_size):
            pages.append("\n\n".join(sections[i: i + chunk_size]))

        return pages, raw_text

    def _clean_text(self, text: str) -> str:
        # """
        # Remove noise from extracted text:
        # - Excessive whitespace and blank lines
        # - Common PDF artifacts (page numbers, headers/footers patterns)
        # - Null bytes and control characters
        # """
        text = re.sub(r'\x00', '', text)
        text = re.sub(r'[\x01-\x08\x0b\x0c\x0e-\x1f\x7f]', '', text)
        text = re.sub(r' +', ' ', text)
        text = re.sub(r'\n{4,}', '\n\n\n', text)
        text = re.sub(r'^\s*\d+\s*$', '', text, flags=re.MULTILINE)
        text = re.sub(r'-{3,}', '---', text)
        text = text.strip()

        return text

    def _detect_title(self, text: str, filename: str) -> Optional[str]:
        """
        Try to extract the document title.
        Strategy: look for patterns common in financial documents.
        Falls back to the filename if nothing is found.
        """
        title_patterns = [
            r'(?i)^(ANNUAL REPORT|10-K|FORM 10-K|QUARTERLY REPORT|10-Q)\b',
            r'(?i)(ANNUAL REPORT\s+(?:FOR|FISCAL YEAR)\s+\d{4})',
            r'(?i)((?:FORM|SEC)\s+10-[KQ])',
        ]

        first_500_chars = text[:500]

        for pattern in title_patterns:
            match = re.search(pattern, first_500_chars)
            if match:
                return match.group(0).strip()

        first_line = text.split('\n')[0].strip()
        if 10 < len(first_line) < 100:
            return first_line

        return Path(filename).stem.replace('_', ' ').replace('-', ' ')

    def _detect_date(self, text: str) -> Optional[str]:
        """
        Find the most prominent date in the document.
        Looks for fiscal year patterns common in financial filings.
        """
        date_patterns = [
            r'(?i)(?:fiscal year|year ended?|for the year)\s+(\w+ \d{1,2},?\s+\d{4})',
            r'(?i)(?:fiscal year|year ended?|for the year)\s+(\d{4})',
            r'(\b(?:January|February|March|April|May|June|July|August|'
            r'September|October|November|December)\s+\d{1,2},\s+\d{4}\b)',
            r'\b(\d{4})\b',
        ]

        for pattern in date_patterns:
            match = re.search(pattern, text[:2000])
            if match:
                return match.group(1).strip()

        return None

    def _detect_category(self, text: str, filename: str) -> str:
        """
        Classify the document type based on content and filename signals.
        Returns a human-readable category string.
        """
        text_lower = text[:3000].lower()
        filename_lower = filename.lower()

        if any(term in text_lower for term in
            ['annual report', '10-k', 'form 10-k', 'fiscal year',
                'shareholders', 'stockholders']):
            return "Annual Report (10-K)"

        if any(term in text_lower for term in
            ['10-q', 'quarterly report', 'quarter ended']):
            return "Quarterly Report (10-Q)"

        if any(term in text_lower for term in
            ['agreement', 'contract', 'whereas', 'hereinafter',
                'party of the first', 'terms and conditions']):
            return "Legal Contract"

        if any(term in text_lower for term in
            ['abstract', 'introduction', 'methodology',
                'conclusion', 'references', 'doi']):
            return "Research Paper"

        if any(term in text_lower for term in
            ['prospectus', 'offering', 'underwriter', 's-1']):
            return "IPO Prospectus"

        return "General Document"